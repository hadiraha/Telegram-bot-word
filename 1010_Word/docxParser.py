from docx import Document

class DocxParser:
    def __init__(self, file_path):
        try:
            self.file_path = file_path
            self.document = Document(file_path)
            print(f"Successfully loaded the document: {file_path}")
        except Exception as e:
            print(f"Error loading document: {e}")
            raise


    def extract_headings_content(self):
        
        # Extract content under Heading 4 and return it in a list.
        # Handles cases where Heading 4 has no text and ensures sections without content are skipped.
        
        content = []
        current_content = None
        collecting = False  # Flag to indicate we are collecting content under a Heading 4
        custom_heading_text = "خبر!"  # Your custom text

        # Print all paragraph styles for debugging
        print("Document Styles Debugging:")
        styles_found = set(para.style.name for para in self.document.paragraphs)
        print(f"Styles found in the document: {styles_found}")
        
        for para in self.document.paragraphs:
            # Check the paragraph style (Heading 4 is what we need)
            if para.style.name == 'Heading 4':
                heading_text = para.text.strip()
                print(f"Detected 'Heading 4': {heading_text}")  # Debugging line
                
                # Store previous content if it has any meaningful content
                if collecting and current_content and current_content.strip() != custom_heading_text:
                    content.append(current_content)
                
                # Start new content collection; use custom text if the heading is empty
                current_content = heading_text if heading_text else custom_heading_text
                collecting = True
            elif collecting:  # Collect content if inside a Heading 4 section
                if para.text.strip():  # Ignore empty paragraphs
                    current_content += "\n" + para.text.strip()

        # Add the last collected content if it has meaningful content
        if collecting and current_content and current_content.strip() != custom_heading_text:
            content.append(current_content)

        return content
    # def extract_headings_content(self):
    #     """
    #     Extract content under Heading 4 and return it in a list.
    #     Handles cases where Heading 4 has no text but content follows in Normal style.
    #     """
    #     content = []
    #     current_content = None
    #     collecting = False  # Flag to indicate we are collecting content under a Heading 4
    #     custom_heading_text = "خبر!"  # Your custom text

    #     # Print all paragraph styles for debugging
    #     print("Document Styles Debugging:")
    #     styles_found = set(para.style.name for para in self.document.paragraphs)
    #     print(f"Styles found in the document: {styles_found}")
        
    #     for para in self.document.paragraphs:
    #         # Check the paragraph style (Heading 4 is what we need)
    #         if para.style.name == 'Heading 4':
    #             heading_text = para.text.strip()
    #             print(f"Detected 'Heading 4': {heading_text}")  # Debugging line

    #             # Store previous content if we were collecting
    #             if collecting and current_content:
    #                 content.append(current_content)
                
    #             # Use custom text if the heading is empty
    #             current_content = heading_text if heading_text else custom_heading_text
    #             collecting = True
    #         elif collecting:  # Collect content if inside a Heading 4 section
    #             current_content += "\n" + para.text.strip()

    #     # Add the last collected content
    #     if collecting and current_content:
    #         content.append(current_content)

    #     return content

# Testing the extraction process
# if __name__ == "__main__":
#     file_path = ''  # Change this to your file path
#     try:
#         docx_parser = DocxParser(file_path)
#         extracted_content = docx_parser.extract_headings_content()
#         if extracted_content:
#             for idx, content in enumerate(extracted_content):
#                 print(f"Content under Heading 4 {idx+1}:")
#                 print(content[:500])  # Print the first 500 characters for preview
#                 print('-' * 50)
#         else:
#             print("No content found under 'Heading 4'. Please check the document.")
#     except Exception as e:
#         print(f"An error occurred: {e}")